import os
from docx import Document
from docx.oxml import parse_xml

def create_nested_hell(output_path):
    doc = Document()
    doc.add_heading('1. Nested Structures Test', level=1)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    a.merge(b)
    a.text = 'Colspan=2'
    
    c = table.cell(1, 0)
    d = table.cell(2, 0)
    c.merge(d)
    c.text = 'Rowspan=2'
    
    target_cell = table.cell(1, 1)
    p = target_cell.add_paragraph('List Item 1 in cell', style='List Bullet')
    p2 = target_cell.add_paragraph('List Item 2 in cell (Nested math: ', style='List Bullet')
    
    math_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>a^2+b^2=c^2</m:t></m:r></m:oMath>'
    p2._element.append(parse_xml(math_xml))
    p2.add_run(')')
    
    target_cell.add_paragraph('[Image: Placeholder for nested image]')
    doc.save(str(output_path))


def create_math_fidelity(output_path):
    doc = Document()
    doc.add_heading('2. Math Fidelity Test', level=1)
    p = doc.add_paragraph("This is an inline OMML formula: ")
    math_xml = '''
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:f>
            <m:num><m:r><m:t>x</m:t></m:r></m:num>
            <m:den><m:r><m:t>y</m:t></m:r></m:den>
        </m:f>
    </m:oMath>
    '''
    p._element.append(parse_xml(math_xml))
    
    doc.add_paragraph("Independent Block Math:")
    math_block = '''
    <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:oMath>
            <m:r><m:t>E = mc^2</m:t></m:r>
        </m:oMath>
    </m:oMathPara>
    '''
    doc._body._element.append(parse_xml(math_block))
    doc.save(str(output_path))


def add_numbered_paragraph(doc, text, ilvl, numId=1):
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    numPr = parse_xml(f'<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:ilvl w:val="{ilvl}"/><w:numId w:val="{numId}"/></w:numPr>')
    pPr.append(numPr)
    return p

def create_multilevel_list(output_path):
    doc = Document()
    doc.add_heading('3. Multi-level List Test', level=1)
    
    add_numbered_paragraph(doc, 'First level 1.', 0)
    add_numbered_paragraph(doc, 'Second level 1.1', 1)
    add_numbered_paragraph(doc, 'Third level a.', 2)
    add_numbered_paragraph(doc, 'Fourth level i.', 3)
    add_numbered_paragraph(doc, 'Second level 1.2', 1)
    add_numbered_paragraph(doc, 'First level 2.', 0)
    
    doc.save(str(output_path))


def create_floating_objects(output_path):
    doc = Document()
    doc.add_heading('4. Floating Objects Test', level=1)
    p = doc.add_paragraph("Here is some text. The next element is a floating textbox in w:drawing/wp:anchor.")
    
    floating_xml = '''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
      <w:r>
        <w:drawing>
          <wp:anchor simplePos="0" relativeHeight="0" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>
            <wp:extent cx="1000000" cy="1000000"/>
            <wp:wrapSquare wrapText="bothSides"/>
            <wp:docPr id="1" name="Textbox 1"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                 <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:v="urn:schemas-microsoft-com:vml">
                   <wps:txbx>
                     <w:txbxContent>
                       <w:p><w:r><w:t>This is a floating text box secretly hiding here.</w:t></w:r></w:p>
                     </w:txbxContent>
                   </wps:txbx>
                 </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </w:r>
    </w:p>
    '''
    # We use append because doc._body._element can append paragraph
    doc._body._element.append(parse_xml(floating_xml))
    
    doc.add_paragraph("Text continues after the floating element.")
    doc.save(str(output_path))


if __name__ == '__main__':
    from pathlib import Path
    out_dir = Path(__file__).parent.parent / "tests" / "data" / "extreme"
    out_dir.mkdir(exist_ok=True, parents=True)
    
    create_nested_hell(out_dir / "1_nested_hell.docx")
    create_math_fidelity(out_dir / "2_math_fidelity.docx")
    create_multilevel_list(out_dir / "3_multilevel_list.docx")
    create_floating_objects(out_dir / "4_floating_objects.docx")
    print(f"✅ Generated 4 extreme docx files in {out_dir}")
