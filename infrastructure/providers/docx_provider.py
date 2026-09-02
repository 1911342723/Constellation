"""
Docx Provider v2 for Constellation — 混合 XML 引擎

混合引擎架构
- 主引擎：python-docx 处理常规段落、表格、图片（稳定可靠）
- 补充引擎：lxml 直接解析 OOXML，捕获 python-docx 忽略的节点：
  - OMML 公式 (w:oMath, w:oMathPara) → Formula Block
  - 浮动文本框 (w:txbxContent) → 递归提取为 Text Block
  - SmartArt / OLE 等未知节点 → [RAW_XML_NODE: tag] 占位符 Block
  
这样既不丢数据，又不用重写整个解析器，兑现"100% 无损率"的承诺。
"""
import hashlib
import logging
import re
from dataclasses import dataclass
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import io
import zipfile

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.document import Document as DocumentType
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.text.run import Run
except ImportError:
    raise ImportError("请安装 python-docx: pip install python-docx") from None

try:
    from lxml import etree
except ImportError:
    etree = None

from app.core.exceptions import ProviderError
from infrastructure.models import Block, StructuralAtom
from modules.parser.heading_candidates import infer_numbering_level

logger = logging.getLogger(__name__)

# OOXML 命名空间
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'v': 'urn:schemas-microsoft-com:vml',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}

ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

# python-docx 能识别的顶级元素类型
KNOWN_ELEMENT_TYPES = (CT_P, CT_Tbl)

# 等宽字体名（小写）→ code 判定；模块级常量避免每次调用重建 set
_MONO_FONTS = frozenset({
    'courier',
    'courier new',
    'consolas',
    'monaco',
    'lucida console',
    'dejavu sans mono',
    'menlo',
    'source code pro',
    'fira code',
    'jetbrains mono',
})

# run 级布尔格式属性，style chain 解析与直接值解析共用同一份清单
_RUN_BOOL_PROPS = ("bold", "italic", "underline", "strike", "superscript", "subscript")

# 已知的无语义 body-level OOXML 元素，静默跳过（纯排版元数据，无可见内容）
SILENT_SKIP_TAGS = frozenset({
    'bookmarkStart', 'bookmarkEnd',
    'proofErr',
    'permStart', 'permEnd',
    'commentRangeStart', 'commentRangeEnd',
    'customXml',
    'moveFromRangeStart', 'moveFromRangeEnd',
    'moveToRangeStart', 'moveToRangeEnd',
    'lastRenderedPageBreak',
    'sectPr',
})

STRICT_OOXML_REPLACEMENTS = (
    (
        b"http://purl.oclc.org/ooxml/officeDocument/relationships/",
        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships/",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/relationships",
        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    ),
    (
        b"http://purl.oclc.org/ooxml/wordprocessingml/main",
        b"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    ),
    (
        b"http://purl.oclc.org/ooxml/drawingml/main",
        b"http://schemas.openxmlformats.org/drawingml/2006/main",
    ),
    (
        b"http://purl.oclc.org/ooxml/officeDocument/math",
        b"http://schemas.openxmlformats.org/officeDocument/2006/math",
    ),
)


# ================================================================
# RichSegment — Run 的标准化中间表示 (Phase 1 output)
# ================================================================

@dataclass(frozen=True, slots=True)
class RichSegment:
    """Immutable, provider-agnostic representation of a single text run.

    Decouples formatting detection (Phase 1) from Markdown rendering
    (Phase 3) so each phase can be tested and evolved independently.
    """

    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    superscript: bool = False
    subscript: bool = False
    code: bool = False
    trusted_markup: bool = False

    @property
    def style_key(self) -> Tuple[bool, ...]:
        """Hashable style fingerprint for homogeneous-run merging."""
        return (self.bold, self.italic, self.underline,
                self.strike, self.superscript, self.subscript, self.code,
                self.trusted_markup)

    @property
    def has_formatting(self) -> bool:
        return any(self.style_key)


@dataclass(slots=True)
class _ResolvedRun:
    """Effective per-run formatting resolved exactly once per run.

    Resolving direct + character-style + paragraph-style values used to be
    the hottest path in DOCX extraction (every property access re-walked
    style chains, and python-docx re-scans the styles part on each
    ``.style`` lookup).  This struct lets the rich-text renderer, the
    block-level detectors and the atom builder share one resolution.
    """

    index: int
    text: str
    bold: bool
    italic: bool
    underline: bool
    strike: bool
    superscript: bool
    subscript: bool
    font_size: Optional[float]
    font_name: Optional[str]
    is_code: bool
    style_id: str


class DocxProvider:
    """
    Docx 文档提供器 v2（混合 XML 引擎）
    
    职责：
    1. 按照文档物理流顺序遍历 .docx 内容
    2. python-docx 处理段落、表格、图片（主引擎）
    3. lxml 补充捕获 OMML 公式、文本框、未知节点（补充引擎）
    4. 嗅探物理特征（Bold/Size/Center/Heading），注入 Block 元数据
    5. 保持零损耗的内容提取
    """
    
    def __init__(self):
        self.image_counter = 0
        self.image_store: Dict[str, bytes] = {}
        self._doc_rels = None
        self._doc_part = None
        self._chain_cache: Dict[tuple, list] = {}
        self._chain_props_cache: Dict[tuple, dict] = {}

    def _reset_state(self) -> None:
        """Reset per-extraction mutable state to avoid cross-document leaks."""
        self.image_counter = 0
        self.image_store = {}
        self._doc_rels = None
        self._doc_part = None
        self._chain_cache = {}
        self._chain_props_cache = {}

    @staticmethod
    def _ensure_supported_path(file_path: str) -> Path:
        path = Path(file_path)
        if path.suffix.lower() != ".docx":
            raise ProviderError("仅支持 .docx 文件；旧版 .doc 暂不支持")
        return path

    @staticmethod
    def _normalize_strict_ooxml(file_bytes: bytes) -> bytes | None:
        source = io.BytesIO(file_bytes)
        output = io.BytesIO()
        replacements_applied = 0

        try:
            with zipfile.ZipFile(source, "r") as input_zip, zipfile.ZipFile(output, "w") as output_zip:
                for zip_info in input_zip.infolist():
                    data = input_zip.read(zip_info.filename)

                    if zip_info.filename.endswith((".xml", ".rels")):
                        updated = data
                        for old, new in STRICT_OOXML_REPLACEMENTS:
                            count = updated.count(old)
                            if count:
                                updated = updated.replace(old, new)
                                replacements_applied += count
                        data = updated

                    output_zip.writestr(zip_info, data)
        except zipfile.BadZipFile:
            return None

        if replacements_applied == 0:
            return None

        return output.getvalue()

    @staticmethod
    def _load_document_from_bytes(file_bytes: bytes):
        if not file_bytes:
            raise ProviderError("上传的 .docx 文件为空")

        try:
            return Document(io.BytesIO(file_bytes))
        except zipfile.BadZipFile as exc:
            raise ProviderError("上传文件不是有效的 .docx 压缩包") from exc
        except (KeyError, ValueError) as exc:
            normalized_bytes = DocxProvider._normalize_strict_ooxml(file_bytes)
            if normalized_bytes is not None:
                logger.info("检测到 Strict Open XML，已转换为 Transitional OOXML 后重试解析")
                try:
                    return Document(io.BytesIO(normalized_bytes))
                except Exception as retry_exc:
                    raise ProviderError(
                        "该 .docx 使用 Strict Open XML 或非标准 OOXML 打包格式，自动兼容转换后仍无法解析；"
                        "请在 Word/WPS 中另存为常规 .docx 后重试"
                    ) from retry_exc

            raise ProviderError(f"无法解析 .docx 文件：{exc}") from exc
    
    def extract(self, file_path: str) -> List[Block]:
        """从文件路径提取 .docx 内容为 Block 列表"""
        self._reset_state()
        path = self._ensure_supported_path(file_path)
        logger.info("开始解析 Docx 文件: %s", path)
        try:
            doc = self._load_document_from_bytes(path.read_bytes())
            blocks = self._extract_blocks(doc)
            logger.info("解析完成，共提取 %d 个 Block", len(blocks))
            self._log_debug_info(blocks)
            return blocks
        except Exception as e:
            logger.error(f"解析 Docx 文件失败: {str(e)}")
            raise
    
    def extract_from_bytes(self, file_bytes: bytes) -> List[Block]:
        """从字节流提取 .docx 内容为 Block 列表"""
        self._reset_state()
        logger.info("开始解析 Docx 字节流")
        try:
            doc = self._load_document_from_bytes(file_bytes)
            blocks = self._extract_blocks(doc)
            logger.info(f"解析完成，共提取 {len(blocks)} 个 Block")
            return blocks
        except Exception as e:
            logger.error(f"解析 Docx 字节流失败: {str(e)}")
            raise
    
    def _extract_blocks(self, doc: DocumentType) -> List[Block]:
        """
        从 Document 对象提取 Block 列表（混合引擎）
        
        遍历 w:body 的每个子元素：
        - CT_P (段落) → python-docx 主引擎处理
        - CT_Tbl (表格) → python-docx 主引擎处理
        - 其他未知元素 → lxml 补充引擎捕获
        """
        blocks: List[Block] = []
        block_id = 0
        self._doc_rels = doc.part.rels
        self._doc_part = doc.part
        self._chain_cache = {}
        self._chain_props_cache = {}
        
        for element_index, element in enumerate(doc.element.body):
            try:
                if isinstance(element, CT_P):
                    # 主引擎：处理段落
                    paragraph = Paragraph(element, doc)
                    
                    # 检查段落内是否包含 OMML 公式
                    formula_blocks = self._extract_omml_from_paragraph(element, block_id)

                    # paragraph.text 是一次 xpath 求值 + 逐 run 拼接，整段流程只算一次
                    para_text = paragraph.text
                    has_plain_text = bool(para_text.strip())
                    has_inline_images = self._paragraph_has_inline_image(element)
                    has_text = has_plain_text or self._paragraph_has_inline_omml(element) or (
                        has_plain_text and has_inline_images
                    )
                    image_blocks = []
                    if not (has_plain_text and has_inline_images):
                        image_blocks = self._extract_images_from_paragraph(
                            paragraph, block_id + (1 if has_text else 0) + len(formula_blocks)
                        )
                    
                    # 检查段落内是否包含文本框
                    textbox_blocks = self._extract_textbox_from_element(
                        element, block_id + (1 if has_text else 0) + len(formula_blocks) + len(image_blocks)
                    )
                    
                    if has_text:
                        text_block = self._process_paragraph(
                            paragraph, block_id, source_index=element_index,
                            raw_text=para_text,
                        )
                        if text_block:
                            blocks.append(text_block)
                            block_id += 1
                    elif self._detect_horizontal_rule(element):
                        # 空段落 + 底部边框 = Word 自动生成的水平分割线
                        blocks.append(Block(
                            id=block_id,
                            type="text",
                            text="---",
                            metadata={"source": "horizontal_rule"},
                        ))
                        block_id += 1
                    
                    for fb in formula_blocks:
                        fb.id = block_id
                        blocks.append(fb)
                        block_id += 1
                    
                    for img_block in image_blocks:
                        img_block.id = block_id
                        blocks.append(img_block)
                        block_id += 1
                    
                    for tb in textbox_blocks:
                        tb.id = block_id
                        blocks.append(tb)
                        block_id += 1
                
                elif isinstance(element, CT_Tbl):
                    # 主引擎：处理表格
                    table = Table(element, doc)
                    table_block = self._process_table(table, block_id)
                    if table_block:
                        blocks.append(table_block)
                        block_id += 1
                        
                elif etree is not None and (element.tag == f'{{{NS["m"]}}}oMathPara' or element.tag == f'{{{NS["m"]}}}oMath'):
                    formula_text = self._omml_to_text(element)
                    if formula_text:
                        blocks.append(Block(
                            id=block_id,
                            type="formula",
                            text=formula_text,
                            metadata={"source": "omml_root"},
                        ))
                        block_id += 1
                
                else:
                    # 前置过滤：静默跳过已知的无语义元数据节点
                    tag_short = element.tag.split('}')[-1] if '}' in str(element.tag) else str(element.tag)
                    if tag_short in SILENT_SKIP_TAGS:
                        logger.debug(f"[混合引擎] 静默跳过已知元数据节点: {tag_short}")
                        continue
                    
                    # 补充引擎：捕获 python-docx 无法识别的节点
                    raw_blocks = self._capture_unknown_element(element, block_id)
                    for rb in raw_blocks:
                        rb.id = block_id
                        blocks.append(rb)
                        block_id += 1
                        
            except Exception as e:
                logger.warning(f"处理元素时出错，安全跳过: {str(e)}")
                continue
        
        return self._post_process_blocks(blocks)
    
    def _post_process_blocks(self, blocks: List[Block]) -> List[Block]:
        """Merge consecutive code blocks without discarding their source atoms."""
        if not blocks:
            return blocks

        merged: List[Block] = []
        for block in blocks:
            if block.type == "code" and merged and merged[-1].type == "code":
                previous = merged[-1]
                previous.text = (previous.text or "") + "\n" + (block.text or "")
                previous_meta = previous.metadata or {}
                previous_meta.setdefault("atoms", []).extend(
                    (block.metadata or {}).get("atoms", [])
                )
                self._recalculate_atom_offsets(previous.text or "", previous_meta["atoms"])
                previous.metadata = previous_meta
                self._apply_atom_weighted_format(previous)
            else:
                merged.append(block)

        # Reassign IDs and update only the lightweight atom dictionaries.
        for i, block in enumerate(merged):
            block.id = i
            for atom in (block.metadata or {}).get("atoms", []):
                atom["block_id"] = i

        return merged

    @staticmethod
    def _recalculate_atom_offsets(block_text: str, atoms: list[dict]) -> None:
        cursor = 0
        for atom in atoms:
            rendered = str(atom.get("provenance", {}).get("canonical_text", atom.get("text", "")))
            start = block_text.find(rendered, cursor) if rendered else cursor
            if start < 0 and rendered.strip():
                trimmed = rendered.strip()
                trimmed_start = block_text.find(trimmed, cursor)
                if trimmed_start >= 0:
                    start = trimmed_start
                    rendered = trimmed
                    atom.setdefault("provenance", {})["canonical_text"] = trimmed
                    atom["provenance"]["join_normalization"] = "paragraph-edge-whitespace-trimmed"
            if start < 0:
                start = min(cursor, len(block_text))
                end = min(len(block_text), start + len(rendered))
                atom.setdefault("provenance", {})["offset_alignment"] = "approximate"
            else:
                end = start + len(rendered)
            atom["char_start"] = start
            atom["char_end"] = end
            cursor = end

    @staticmethod
    def _visible_char_weight(text: str) -> int:
        return sum(1 for char in (text or "") if not char.isspace())

    @classmethod
    def _apply_atom_weighted_format(cls, block: Block) -> None:
        from collections import Counter

        atoms = (block.metadata or {}).get("atoms", [])
        total_weight = 0
        bold_weight = 0
        sizes: Counter[float] = Counter()
        for atom in atoms:
            weight = cls._visible_char_weight(str(atom.get("text", "")))
            if not weight:
                continue
            total_weight += weight
            if atom.get("is_bold") is True:
                bold_weight += weight
            if atom.get("font_size"):
                sizes[round(float(atom["font_size"]), 1)] += weight
        if total_weight:
            block.is_bold = bold_weight >= total_weight * 0.8
        if sizes:
            block.font_size = sizes.most_common(1)[0][0]
    
    # ================================================================
    # v2 新增：OMML 公式提取
    # ================================================================
    
    def _extract_omml_from_paragraph(self, element, start_block_id: int) -> List[Block]:
        """
        从段落 XML 中提取 OMML 公式 (w:oMath / w:oMathPara)
        
        python-docx 会忽略公式节点，这里用 lxml 直接捕获。
        """
        if etree is None:
            return []
        
        formula_blocks = []
        bid = start_block_id
        
        # 查找段落内的 oMathPara（独立公式块）和 oMath（行内公式）
        for math_elem in element.findall(f'.//{{{NS["m"]}}}oMathPara'):
            formula_text = self._omml_to_text(math_elem)
            if formula_text:
                formula_blocks.append(Block(
                    id=bid,
                    type="formula",
                    text=formula_text,
                    metadata={"source": "omml_para"},
                ))
                bid += 1
        
        return formula_blocks

    def _paragraph_has_inline_omml(self, element) -> bool:
        if etree is None:
            return False

        for math_elem in element.findall(f'.//{{{NS["m"]}}}oMath'):
            parent = math_elem.getparent()
            if parent is not None and parent.tag == f'{{{NS["m"]}}}oMathPara':
                continue
            return True
        return False

    def _paragraph_has_inline_image(self, element) -> bool:
        return element.find(f'.//{{{NS["w"]}}}drawing') is not None or element.find(f'.//{{{NS["w"]}}}pict') is not None
    
    def _omml_to_text(self, element) -> str:
        """把 OMML (Office MathML) 元素降级为 LaTeX 风格的可读文本。

        覆盖常见结构，避免复杂公式被"通用遍历"拍平成无意义字符串
        （如 x² 退化成 x2、√x 退化成 x、Σ 上下限丢失）：

        - ``m:f``    分数        → ``\\frac{num}{den}``
        - ``m:sSup`` / ``m:sSub`` / ``m:sSubSup`` 上下标 → ``base^{..}`` / ``base_{..}``
        - ``m:rad``  根式        → ``\\sqrt[deg]{e}`` / ``\\sqrt{e}``
        - ``m:nary`` 求和 / 积分 → ``op_{sub}^{sup} e``（运算符取 m:chr，默认 ∫）
        - ``m:d``    定界符      → ``(e)``（括号取 m:begChr / m:endChr）
        - ``m:func`` 函数        → ``fName(e)``
        - ``m:limLow`` / ``m:limUpp`` 上下极限 → ``base_{lim}`` / ``base^{lim}``
        - ``m:r``    文本 run    → 拼接其中所有 m:t
        - 其它容器（oMath / e / num / sub ...）→ 递归拼接子节点

        遇到任何异常都安全降级为通用遍历，绝不让单个公式拖垮整篇解析。
        """
        if element is None:
            return ""

        m = NS["m"]
        raw_tag = element.tag
        tag = raw_tag.split('}', 1)[-1] if isinstance(raw_tag, str) else ""

        def _text_of(name: str) -> str:
            node = element.find(f'{{{m}}}{name}')
            return self._omml_to_text(node) if node is not None else ""

        def _attr_val(path: str, default: str) -> str:
            node = element.find(path)
            if node is None:
                return default
            val = node.get(f'{{{m}}}val')
            return val if val else default

        try:
            if tag == "f":
                return f"\\frac{{{_text_of('num')}}}{{{_text_of('den')}}}"
            if tag == "r":
                return "".join(t.text or "" for t in element.findall(f'{{{m}}}t'))
            if tag == "sSup":
                return f"{_text_of('e')}^{{{_text_of('sup')}}}"
            if tag == "sSub":
                return f"{_text_of('e')}_{{{_text_of('sub')}}}"
            if tag == "sSubSup":
                return f"{_text_of('e')}_{{{_text_of('sub')}}}^{{{_text_of('sup')}}}"
            if tag == "rad":
                deg, body = _text_of('deg'), _text_of('e')
                return f"\\sqrt[{deg}]{{{body}}}" if deg else f"\\sqrt{{{body}}}"
            if tag == "nary":
                op = _attr_val(f'{{{m}}}naryPr/{{{m}}}chr', "∫") or "∫"
                sub, sup, body = _text_of('sub'), _text_of('sup'), _text_of('e')
                out = op
                if sub:
                    out += f"_{{{sub}}}"
                if sup:
                    out += f"^{{{sup}}}"
                if body:
                    out += f" {body}"
                return out
            if tag == "d":
                beg = _attr_val(f'{{{m}}}dPr/{{{m}}}begChr', "(")
                end = _attr_val(f'{{{m}}}dPr/{{{m}}}endChr', ")")
                inner = "".join(self._omml_to_text(e) for e in element.findall(f'{{{m}}}e'))
                return f"{beg}{inner}{end}"
            if tag == "func":
                return f"{_text_of('fName')}({_text_of('e')})"
            if tag == "limLow":
                return f"{_text_of('e')}_{{{_text_of('lim')}}}"
            if tag == "limUpp":
                return f"{_text_of('e')}^{{{_text_of('lim')}}}"
        except Exception as exc:  # noqa: BLE001 - 公式结构异常时降级为通用遍历
            logger.debug("[OMML] 结构化解析失败，降级为通用遍历: %s", exc)

        # 通用遍历（容器节点 / 未覆盖结构 / 异常降级）
        texts = []
        for child in element:
            t = self._omml_to_text(child)
            if t:
                texts.append(t)
        return "".join(texts).strip()
    
    # ================================================================
    # v2 新增：浮动文本框提取
    # ================================================================
    
    def _extract_textbox_from_element(self, element, start_block_id: int) -> List[Block]:
        """
        从元素中提取浮动文本框 (w:txbxContent) 的内容
        
        文本框内容在 python-docx 中被忽略，这里递归提取。
        """
        if etree is None:
            return []
        
        textbox_blocks = []
        bid = start_block_id
        
        # 搜索所有 txbxContent 节点
        for txbx in element.iter(f'{{{NS["w"]}}}txbxContent'):
            # 递归提取文本框内的段落文本
            paragraphs_text = []
            for p_elem in txbx.findall(f'{{{NS["w"]}}}p'):
                p_text = self._extract_text_from_xml_paragraph(p_elem)
                if p_text.strip():
                    paragraphs_text.append(p_text.strip())
            
            if paragraphs_text:
                combined = "\n".join(paragraphs_text)
                textbox_blocks.append(Block(
                    id=bid,
                    type="text",
                    text=combined,
                    metadata={"source": "textbox"},
                ))
                bid += 1
        
        return textbox_blocks
    
    def _extract_text_from_xml_paragraph(self, p_element) -> str:
        """从 XML 段落元素中提取纯文本"""
        texts = []
        for t_elem in p_element.iter(f'{{{NS["w"]}}}t'):
            if t_elem.text:
                texts.append(t_elem.text)
        return "".join(texts)
    
    # ================================================================
    # v2 新增：未知节点捕获
    # ================================================================
    
    def _capture_unknown_element(self, element, block_id: int) -> List[Block]:
        """
        捕获 python-docx 无法识别的 XML 节点
        
        对于 SmartArt、OLE、自定义 XML 等复杂节点，
        记录其标签名和内部文本作为占位符 Block，
        确保即便我们不理解这是什么，最终组装时也能原封不动地保留位置。
        """
        tag = element.tag if hasattr(element, 'tag') else str(type(element))
        
        # 清理命名空间前缀，提取可读标签名
        tag_short = tag.split('}')[-1] if '}' in tag else tag
        
        # 尝试提取内部文本
        inner_text = ""
        try:
            if etree is not None:
                # 提取所有 w:t 文本节点
                t_texts = []
                for t_elem in element.iter(f'{{{NS["w"]}}}t'):
                    if t_elem.text:
                        t_texts.append(t_elem.text)
                inner_text = "".join(t_texts).strip()
        except Exception:
            pass
        
        if not inner_text:
            # 如果没有文本内容，记录为纯占位符
            placeholder = f"[RAW_XML_NODE: {tag_short}]"
        else:
            placeholder = f"[RAW_XML_NODE: {tag_short}] {inner_text}"
        
        logger.debug(f"[混合引擎] 捕获未知节点: {tag_short}, 内容长度: {len(inner_text)}")
        
        return [Block(
            id=block_id,
            type="text",
            text=placeholder,
            metadata={"source": "raw_xml", "tag": tag_short},
        )]
    
    # ================================================================
    # 原有逻辑（保持不变）
    # ================================================================
    
    @staticmethod
    def _read_w_val(element, tag: str) -> Optional[str]:
        """Read a ``w:val`` attribute from the first child named *tag*."""
        if element is None:
            return None
        child = element.find(f'{{{NS["w"]}}}{tag}')
        if child is None:
            return None
        return child.get(f'{{{NS["w"]}}}val')

    @staticmethod
    def _style_chain(style) -> list:
        """Return closest-first basedOn chain with cycle protection."""
        chain = []
        seen: set[str] = set()
        current = style
        while current is not None:
            key = str(getattr(current, "style_id", "") or id(current))
            if key in seen:
                break
            seen.add(key)
            chain.append(current)
            try:
                current = current.base_style
            except Exception:
                break
        return chain

    # ================================================================
    # 缓存化有效格式解析
    #
    # python-docx 的 paragraph.style / run.style 访问远非 O(1)：无显式
    # style id 时每次都对 styles part 做线性 XML 扫描（Styles.default），
    # 旧实现"每 run × 每属性重建 style chain"把该成本放大成 90%+ 的
    # 提取耗时。这里按 (style_id, style_type) 缓存 basedOn 链与"链上
    # 首个非 None 字体属性"，每文档每样式只解析一次。
    # ================================================================

    def _part_for(self, owner) -> Optional[object]:
        """Return the document part, deriving it lazily for direct calls."""
        if self._doc_part is None:
            try:
                self._doc_part = owner.part
            except Exception:
                return None
        return self._doc_part

    def _chain_for_style_id(self, style_id: Optional[str], style_type) -> list:
        key = (style_id, style_type)
        chain = self._chain_cache.get(key)
        if chain is None:
            style = None
            if self._doc_part is not None:
                try:
                    style = self._doc_part.get_style(style_id, style_type)
                except Exception:
                    style = None
            chain = self._style_chain(style)
            self._chain_cache[key] = chain
        return chain

    def _paragraph_chain(self, paragraph: Paragraph) -> list:
        """Cached closest-first basedOn chain for a paragraph's style."""
        if self._part_for(paragraph) is None:
            return self._style_chain(paragraph.style)
        try:
            style_id = paragraph._p.style
        except Exception:
            style_id = None
        return self._chain_for_style_id(style_id, WD_STYLE_TYPE.PARAGRAPH)

    def _chain_props(self, style_id: Optional[str], style_type) -> dict:
        """First non-None font property along one basedOn chain (cached)."""
        key = (style_id, style_type)
        props = self._chain_props_cache.get(key)
        if props is None:
            props = self._props_from_chain(self._chain_for_style_id(style_id, style_type))
            self._chain_props_cache[key] = props
        return props

    @staticmethod
    def _props_from_chain(chain: list) -> dict:
        props: dict = {name: None for name in _RUN_BOOL_PROPS}
        props["size_pt"] = None
        props["name"] = None
        for style in chain:
            try:
                font = style.font
            except Exception:
                continue
            for name in _RUN_BOOL_PROPS:
                if props[name] is None:
                    try:
                        value = getattr(font, name)
                    except Exception:
                        continue
                    if value is not None:
                        props[name] = bool(value)
            if props["size_pt"] is None:
                try:
                    if font.size is not None:
                        props["size_pt"] = float(font.size.pt)
                except Exception:
                    pass
            if props["name"] is None:
                try:
                    if font.name:
                        props["name"] = str(font.name)
                except Exception:
                    pass
        return props

    def _resolve_paragraph_runs(self, paragraph: Paragraph) -> tuple[list[_ResolvedRun], bool]:
        """Resolve every non-empty run of a paragraph exactly once.

        Returns ``(resolved_runs, had_any_runs)``; the flag preserves the
        legacy distinction between "no runs at all" (rich text falls back
        to ``paragraph.text``) and "runs exist but are all empty".
        """
        runs = paragraph.runs
        # run.text 是一次 xpath 求值，读一遍后直接传入解析，避免重复求值。
        resolved = [
            self._resolve_run(run, index, paragraph, text)
            for index, (run, text) in enumerate((run, run.text) for run in runs)
            if text
        ]
        return resolved, bool(runs)

    def _resolve_run(self, run, index: int, paragraph: Paragraph, text: str) -> _ResolvedRun:
        if self._part_for(paragraph) is not None:
            try:
                run_style_id = run._r.style
            except Exception:
                run_style_id = None
            try:
                para_style_id = paragraph._p.style
            except Exception:
                para_style_id = None
            run_chain = self._chain_for_style_id(run_style_id, WD_STYLE_TYPE.CHARACTER)
            run_props = self._chain_props(run_style_id, WD_STYLE_TYPE.CHARACTER)
            para_props = self._chain_props(para_style_id, WD_STYLE_TYPE.PARAGRAPH)
        else:
            run_chain = self._style_chain(getattr(run, "style", None))
            run_props = self._props_from_chain(run_chain)
            para_props = self._props_from_chain(self._style_chain(paragraph.style))

        # run 的全部直接格式值（布尔属性 / 字号 / 字体名）都存放在 rPr 里；
        # 无 rPr 的纯文本 run（多数文档的主体）可跳过所有直接值探测。
        try:
            has_rpr = run._r.rPr is not None
        except Exception:
            has_rpr = True

        def effective_bool(name: str) -> bool:
            direct = self._direct_font_value(run, name) if has_rpr else None
            if direct is not None:
                return bool(direct)
            value = run_props[name]
            if value is None:
                value = para_props[name]
            return bool(value) if value is not None else False

        font_size: Optional[float] = None
        if has_rpr:
            try:
                if run.font.size is not None:
                    font_size = float(run.font.size.pt)
            except Exception:
                pass
        if font_size is None:
            font_size = (
                run_props["size_pt"]
                if run_props["size_pt"] is not None
                else para_props["size_pt"]
            )

        font_name = self._get_run_font_name(run) if has_rpr else None
        if not font_name:
            font_name = run_props["name"] or para_props["name"]

        resolved_style = run_chain[0] if run_chain else None

        return _ResolvedRun(
            index=index,
            text=text,
            bold=effective_bool("bold"),
            italic=effective_bool("italic"),
            underline=effective_bool("underline"),
            strike=effective_bool("strike"),
            superscript=effective_bool("superscript"),
            subscript=effective_bool("subscript"),
            font_size=font_size,
            font_name=font_name,
            is_code=bool(font_name and font_name.lower() in _MONO_FONTS),
            style_id=str(getattr(resolved_style, "style_id", "") or ""),
        )

    def _paragraph_outline_level(self, paragraph: Paragraph) -> Optional[int]:
        """Read direct or effective basedOn-chain outline level (0-based)."""
        try:
            value = self._read_w_val(paragraph._element.pPr, "outlineLvl")
            if value is not None:
                return int(value)
        except Exception:
            pass

        for style in self._paragraph_chain(paragraph):
            try:
                value = self._read_w_val(style.element.pPr, "outlineLvl")
                if value is not None:
                    return int(value)
            except Exception:
                continue
        return None

    def _paragraph_list_level(self, paragraph: Paragraph) -> Optional[int]:
        """Read effective direct/style basedOn-chain ``numPr`` (0-based).

        OOXML permits ``numPr`` with a ``numId`` but no explicit ``ilvl``;
        that form denotes level zero.  A direct ``numId=0`` explicitly removes
        numbering and must stop inherited style numbering from leaking through.
        """
        property_sources = []
        try:
            property_sources.append(paragraph._element.pPr)
        except Exception:
            pass
        for style in self._paragraph_chain(paragraph):
            try:
                property_sources.append(style.element.pPr)
            except Exception:
                continue

        for p_pr in property_sources:
            if p_pr is None:
                continue
            try:
                num_pr = p_pr.find(f'{{{NS["w"]}}}numPr')
                if num_pr is None:
                    continue
                num_id = self._read_w_val(num_pr, "numId")
                if num_id == "0":
                    return None
                level = self._read_w_val(num_pr, "ilvl")
                if level is not None:
                    return int(level)
                if num_id is not None:
                    return 0
            except (TypeError, ValueError):
                continue
        return None

    def _effective_style_identity(self, paragraph: Paragraph) -> tuple[str, str, list[str]]:
        """Return the style that contributes heading structure plus full chain."""
        chain = self._paragraph_chain(paragraph)
        chain_ids = [str(getattr(style, "style_id", "") or "") for style in chain]
        heading_pattern = re.compile(
            r"^(?:Heading\s*|标题\s*|標題\s*|Title\s*)(\d+)$",
            re.IGNORECASE,
        )
        for style in chain:
            name = str(getattr(style, "name", "") or "")
            style_id = str(getattr(style, "style_id", "") or "")
            try:
                has_outline = self._read_w_val(style.element.pPr, "outlineLvl") is not None
            except Exception:
                has_outline = False
            if heading_pattern.match(name) or re.match(r"^Heading\d+$", style_id, re.IGNORECASE) or has_outline:
                return name, style_id, chain_ids
        if chain:
            return (
                str(getattr(chain[0], "name", "") or ""),
                str(getattr(chain[0], "style_id", "") or ""),
                chain_ids,
            )
        return "", "", chain_ids

    def _detect_heading_structure(
        self,
        paragraph: Paragraph,
        text: str,
    ) -> tuple[bool, Optional[int], str, str, Optional[int], Optional[int]]:
        """Detect localized/effective Word heading style and structural levels."""
        chain = self._paragraph_chain(paragraph)
        style = chain[0] if chain else None
        style_name = style.name if style else ""
        style_id = getattr(style, "style_id", "") if style else ""

        outline_level = self._paragraph_outline_level(paragraph)
        list_level = self._paragraph_list_level(paragraph)
        visible_numbering_level = infer_numbering_level(text)
        numbering_level = (
            list_level + 1 if list_level is not None else visible_numbering_level
        )

        heading_level: Optional[int] = None
        patterns = (
            r"^Heading\s*(\d+)$",
            r"^标题\s*(\d+)$",
            r"^標題\s*(\d+)$",
            r"^Heading(\d+)$",
            r"^Title\s*(\d+)$",
        )
        for candidate_style in chain:
            candidate_name = str(getattr(candidate_style, "name", "") or "")
            candidate_id = str(getattr(candidate_style, "style_id", "") or "")
            for pattern in patterns:
                match = re.match(pattern, candidate_name, re.IGNORECASE)
                if match:
                    heading_level = int(match.group(1))
                    break
            if heading_level is None:
                match = re.match(r"^Heading(\d+)$", candidate_id, re.IGNORECASE)
                if match:
                    heading_level = int(match.group(1))
            if heading_level is not None:
                break

        if heading_level is None and outline_level is not None:
            heading_level = outline_level + 1

        is_heading = heading_level is not None
        if heading_level is not None:
            heading_level = max(1, min(heading_level, 6))

        return (
            is_heading,
            heading_level,
            style_name,
            style_id,
            outline_level,
            numbering_level,
        )

    def _process_paragraph(
        self,
        paragraph: Paragraph,
        block_id: int,
        source_index: Optional[int] = None,
        raw_text: Optional[str] = None,
    ) -> Optional[Block]:
        """Process a paragraph while retaining run-level physical atoms.

        ``raw_text`` lets ``_extract_blocks`` share its already-computed
        ``paragraph.text``; direct callers omit it.
        """
        prepared = self._resolve_paragraph_runs(paragraph)
        resolved_runs, _ = prepared
        rich_text = self._extract_rich_text(paragraph, prepared)
        plain_text = (paragraph.text if raw_text is None else raw_text).strip()
        text = rich_text.strip() if rich_text and rich_text.strip() else plain_text
        if not text:
            return None

        (
            is_heading,
            heading_level,
            style_name,
            style_id,
            outline_level,
            numbering_level,
        ) = self._detect_heading_structure(paragraph, plain_text)

        is_bold = self._detect_bold(resolved_runs)
        is_code = self._detect_code_font(resolved_runs)
        if is_code and plain_text:
            text = plain_text
        font_size = self._detect_font_size(paragraph, resolved_runs)
        alignment = self._detect_alignment(paragraph)
        list_level = self._paragraph_list_level(paragraph)
        effective_style_name, effective_style_id, style_chain = (
            self._effective_style_identity(paragraph)
        )
        paragraph_index = block_id if source_index is None else source_index
        atoms = self._build_docx_run_atoms(
            paragraph_index=paragraph_index,
            block_id=block_id,
            block_text=text,
            alignment=alignment,
            effective_style_id=effective_style_id,
            resolved_runs=resolved_runs,
        )

        block = Block(
            id=block_id,
            type="code" if is_code else "text",
            text=text,
            metadata={
                "style": style_name,
                "style_id": style_id,
                "effective_style": effective_style_name,
                "effective_style_id": effective_style_id,
                "style_chain": style_chain,
                "outline_level": outline_level,
                "list_level": list_level,
                "numbering_level": numbering_level,
                "source": "paragraph",
                "source_span": {"paragraph": paragraph_index},
                "atoms": atoms,
                "provenance": {
                    "provider": "python-docx",
                    "effective_properties": "direct+run-style+paragraph-basedOn",
                },
            },
            is_bold=is_bold,
            font_size=font_size,
            alignment=alignment,
            is_heading_style=is_heading,
            heading_level=heading_level,
            has_heading_numbering=numbering_level is not None,
        )
        self._apply_atom_weighted_format(block)
        return block
    
    @staticmethod
    def _direct_font_value(run, property_name: str):
        try:
            if property_name in {"bold", "italic", "underline"}:
                return getattr(run, property_name)
            if property_name in {"superscript", "subscript"}:
                value = getattr(run.font, property_name)
                if value is not None:
                    return bool(value)
                vert_align = run.font._element.find(f'{{{NS["w"]}}}vertAlign')
                if vert_align is not None:
                    val = vert_align.get(f'{{{NS["w"]}}}val')
                    return val == property_name
            return getattr(run.font, property_name)
        except Exception:
            return None

    def _build_docx_run_atoms(
        self,
        *,
        paragraph_index: int,
        block_id: int,
        block_text: str,
        alignment: Optional[str],
        effective_style_id: str,
        resolved_runs: list[_ResolvedRun],
    ) -> list[dict]:
        atoms: list[dict] = []
        for rr in resolved_runs:
            source_span = {"paragraph": paragraph_index, "run": rr.index}
            atom = StructuralAtom.create(
                source="docx_run",
                source_span=source_span,
                block_id=block_id,
                text=rr.text,
                font_family=rr.font_name,
                font_size=rr.font_size,
                is_bold=rr.bold,
                is_italic=rr.italic,
                is_underline=rr.underline,
                is_strike=rr.strike,
                is_superscript=rr.superscript,
                is_subscript=rr.subscript,
                is_code=rr.is_code,
                alignment=alignment,
                region="body",
                provenance={
                    "provider": "python-docx",
                    "run_style_id": rr.style_id,
                    "effective_paragraph_style_id": effective_style_id,
                },
            ).to_metadata()
            atoms.append(atom)

        if not atoms:
            source_span = {"paragraph": paragraph_index}
            atoms.append(StructuralAtom.create(
                source="docx_paragraph",
                source_span=source_span,
                block_id=block_id,
                text=block_text,
                alignment=alignment,
                region="body",
                provenance={
                    "provider": "python-docx",
                    "effective_paragraph_style_id": effective_style_id,
                },
            ).to_metadata())
        self._recalculate_atom_offsets(block_text, atoms)
        return atoms

    # ================================================================
    # v3: Three-Phase Rich Text Pipeline
    #
    #   Resolved Runs → [Phase 1: Normalize] → [Phase 2: Merge] → [Phase 3: Render]
    #
    # Phase 1 (_resolve_paragraph_runs + _segment_from_resolved): Convert
    #   python-docx Runs into immutable RichSegment objects via the shared
    #   cached effective-format resolution.  No string manipulation.
    # Phase 2 (_merge_homogeneous_segments): Absorb adjacent segments
    #   with identical style fingerprints into one.
    # Phase 3 (_render_segments): Apply Strip-Safe Markdown wrapping
    #   to each merged segment and concatenate.
    # ================================================================

    def _extract_rich_text(
        self,
        paragraph: Paragraph,
        prepared: Optional[tuple[list[_ResolvedRun], bool]] = None,
    ) -> str:
        """Three-phase rich text pipeline orchestrator.

        Returns a Markdown string with clean inline formatting:
        no ghost spaces, no fragmented bold runs, no marker bleed.
        ``prepared`` lets ``_process_paragraph`` share its run resolution;
        other callers (table cells) resolve on demand.
        """
        if self._paragraph_has_inline_omml(paragraph._element) or self._paragraph_has_inline_image(paragraph._element):
            return self._extract_rich_text_with_inline_objects(paragraph)

        if prepared is None:
            prepared = self._resolve_paragraph_runs(paragraph)
        resolved_runs, has_runs = prepared
        if not has_runs:
            return paragraph.text or ""

        # Phase 1 → Phase 2 → Phase 3
        segments = [self._segment_from_resolved(rr) for rr in resolved_runs]
        segments = self._merge_homogeneous_segments(segments)
        return self._render_segments(segments)

    @staticmethod
    def _segment_from_resolved(rr: _ResolvedRun) -> RichSegment:
        return RichSegment(
            text=rr.text,
            bold=rr.bold,
            italic=rr.italic,
            underline=rr.underline,
            strike=rr.strike,
            superscript=rr.superscript,
            subscript=rr.subscript,
            code=rr.is_code,
        )

    def _extract_rich_text_with_inline_objects(self, paragraph: Paragraph) -> str:
        segments: List[RichSegment] = []

        for child in paragraph._element:
            if child.tag == f'{{{NS["w"]}}}r':
                segments.extend(self._normalize_run_element_in_order(child, paragraph))
            elif child.tag == f'{{{NS["m"]}}}oMath':
                formula_text = self._omml_to_text(child)
                if formula_text:
                    segments.append(RichSegment(
                        text=Block.render_inline_formula(formula_text), trusted_markup=True,
                    ))
            elif child.tag == f'{{{NS["m"]}}}oMathPara':
                formula_text = self._omml_to_text(child)
                if formula_text:
                    segments.append(RichSegment(
                        text=f"\n{Block.render_block_formula(formula_text)}\n",
                        trusted_markup=True,
                    ))

        segments = self._merge_homogeneous_segments(segments)
        return self._render_segments(segments)

    def _normalize_run_element_in_order(self, run_element, paragraph: Paragraph) -> List[RichSegment]:
        run = Run(run_element, paragraph)
        run_text = run.text
        base_segments = (
            [self._segment_from_resolved(self._resolve_run(run, -1, paragraph, run_text))]
            if run_text else []
        )
        if not self._run_has_inline_image(run_element):
            return base_segments

        segments: List[RichSegment] = []
        text_segments = iter(base_segments)
        for child in run_element:
            if child.tag == f'{{{NS["w"]}}}t':
                try:
                    segments.append(next(text_segments))
                except StopIteration:
                    if child.text:
                        segments.append(RichSegment(text=child.text))
            elif child.tag == f'{{{NS["w"]}}}drawing':
                image_block = self._extract_image_from_drawing(child, -1)
                if image_block:
                    segments.append(RichSegment(text=image_block.to_markdown(), trusted_markup=True))
            elif child.tag == f'{{{NS["w"]}}}pict':
                image_block = self._extract_image_from_pict(child, -1)
                if image_block:
                    segments.append(RichSegment(text=image_block.to_markdown(), trusted_markup=True))

        segments.extend(text_segments)
        return segments

    def _run_has_inline_image(self, run_element) -> bool:
        return run_element.find(f'.//{{{NS["w"]}}}drawing') is not None or run_element.find(f'.//{{{NS["w"]}}}pict') is not None

    # -- Phase 2: Homogeneous-Run Merge ------------------------------------

    @staticmethod
    def _merge_homogeneous_segments(segments: List[RichSegment]) -> List[RichSegment]:
        """Absorb adjacent segments with identical style fingerprints.

        ``**A** **B** **C**`` becomes ``**A B C**``.
        Pure-whitespace segments between two same-style segments are
        absorbed into the preceding segment to prevent orphaned spaces.
        """
        if not segments:
            return segments

        merged: List[RichSegment] = [segments[0]]

        for seg in segments[1:]:
            prev = merged[-1]

            if seg.style_key == prev.style_key:
                # Same style → merge text
                merged[-1] = RichSegment(
                    text=prev.text + seg.text,
                    bold=prev.bold,
                    italic=prev.italic,
                    underline=prev.underline,
                    strike=prev.strike,
                    superscript=prev.superscript,
                    subscript=prev.subscript,
                    code=prev.code,
                    trusted_markup=prev.trusted_markup,
                )
            elif not seg.text.strip() and not seg.has_formatting:
                # Pure whitespace with no formatting — absorb into prev
                merged[-1] = RichSegment(
                    text=prev.text + seg.text,
                    bold=prev.bold,
                    italic=prev.italic,
                    underline=prev.underline,
                    strike=prev.strike,
                    superscript=prev.superscript,
                    subscript=prev.subscript,
                    code=prev.code,
                    trusted_markup=prev.trusted_markup,
                )
            else:
                merged.append(seg)

        return merged

    # -- Phase 3: Strip-Safe Render ----------------------------------------

    @staticmethod
    def _wrap_safe(text: str, marker: str) -> str:
        """Wrap *text* with *marker*, keeping whitespace outside.

        ``_wrap_safe(" hello ", "**")`` → ``" **hello** "``

        Prevents the ghost-space problem where Markdown markers
        bleed into surrounding whitespace.
        """
        if not text or not text.strip():
            return text  # pure whitespace — never wrap

        prefix = text[:len(text) - len(text.lstrip())]
        suffix = text[len(text.rstrip()):]
        core = text.strip()
        return f"{prefix}{marker}{core}{marker}{suffix}"

    @staticmethod
    def _escape_inline_text(text: str) -> str:
        """Escape source characters that could become Markdown/HTML syntax."""
        return re.sub(r"([\\`*_[\]~<>|])", r"\\\1", text)

    @classmethod
    def _render_segments(cls, segments: List[RichSegment]) -> str:
        """Render merged RichSegments into a Markdown string.

        Applies markers inside-out (strike → italic → bold) with
        Strip-Safe wrapping at each layer.
        """
        parts: List[str] = []

        for seg in segments:
            text = seg.text

            if seg.trusted_markup:
                parts.append(text)
                continue

            if seg.code:
                text = text.replace("`", "\\`")
            else:
                text = cls._escape_inline_text(text)

            if not seg.has_formatting:
                parts.append(text)
                continue

            # Apply markers inside-out so nesting is correct
            if seg.code:
                text = cls._wrap_safe(text, "`")
            if seg.strike:
                text = cls._wrap_safe(text, "~~")
            if seg.italic:
                text = cls._wrap_safe(text, "*")
            if seg.bold:
                text = cls._wrap_safe(text, "**")
            if seg.underline:
                stripped = text.strip()
                if stripped:
                    prefix = text[:len(text) - len(text.lstrip())]
                    suffix = text[len(text.rstrip()):]
                    text = f"{prefix}<u>{stripped}</u>{suffix}"
            if seg.superscript:
                text = f"<sup>{text.strip()}</sup>"
            if seg.subscript:
                text = f"<sub>{text.strip()}</sub>"

            parts.append(text)

        return "".join(parts)

    # -- Bold detection (v3: stricter threshold) ---------------------------

    def _detect_bold(self, resolved_runs: list[_ResolvedRun]) -> bool:
        """Detect paragraph boldness by effective non-whitespace character mass."""
        total_weight = 0
        bold_weight = 0
        for rr in resolved_runs:
            weight = self._visible_char_weight(rr.text)
            if not weight:
                continue
            total_weight += weight
            if rr.bold:
                bold_weight += weight
        return total_weight > 0 and bold_weight >= total_weight * 0.8

    def _detect_code_font(self, resolved_runs: list[_ResolvedRun]) -> bool:
        """Detect code paragraphs by effective monospace character mass (>=80%)."""
        total_weight = 0
        code_weight = 0
        for rr in resolved_runs:
            weight = self._visible_char_weight(rr.text)
            if not weight:
                continue
            total_weight += weight
            if rr.is_code:
                code_weight += weight
        return total_weight > 0 and code_weight >= total_weight * 0.8

    @staticmethod
    def _mono_fonts() -> frozenset[str]:
        return _MONO_FONTS

    def _get_run_font_name(self, run) -> Optional[str]:
        if run.font and run.font.name:
            return run.font.name

        try:
            r_fonts = run._element.rPr.rFonts
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                value = getattr(r_fonts, attr, None)
                if value:
                    return value
        except Exception:
            pass

        return None
    
    def _detect_font_size(
        self,
        paragraph: Paragraph,
        resolved_runs: list[_ResolvedRun],
    ) -> Optional[float]:
        """Return the effective font size carrying the most visible characters."""
        from collections import Counter

        sizes: Counter[float] = Counter()
        for rr in resolved_runs:
            weight = self._visible_char_weight(rr.text)
            if not weight:
                continue
            if rr.font_size is not None and rr.font_size > 0:
                sizes[round(rr.font_size, 1)] += weight
        if sizes:
            return sizes.most_common(1)[0][0]

        for style in self._paragraph_chain(paragraph):
            try:
                if style.font.size is not None:
                    return float(style.font.size.pt)
            except Exception:
                continue
        return None
    
    def _detect_horizontal_rule(self, paragraph_element) -> bool:
        """检测段落是否是 Word 自动生成的水平分割线。

        Word 中输入 --- 回车生成的分割线本质上是一个空段落，
        其段落属性中带有底部边框 (w:pBdr > w:bottom)。
        """
        pPr = paragraph_element.find(f'{{{NS["w"]}}}pPr')
        if pPr is None:
            return False
        pBdr = pPr.find(f'{{{NS["w"]}}}pBdr')
        if pBdr is None:
            return False
        bottom = pBdr.find(f'{{{NS["w"]}}}bottom')
        return bottom is not None

    def _detect_alignment(self, paragraph: Paragraph) -> Optional[str]:
        """嗅探段落对齐方式"""
        try:
            if paragraph.alignment is not None:
                return ALIGNMENT_MAP.get(paragraph.alignment, None)
        except Exception:
            pass
        return None
    
    def _process_table(self, table: Table, block_id: int) -> Optional[Block]:
        """处理表格元素，转换为 Markdown 格式并提取完整的合并单元格信息。"""
        try:
            rows = []
            cells_info = []
            # python-docx 的 row.cells 每次调用都重建整张表的 cell 网格
            # （Table._cells 全表扫描），逐行访问是 O(行²)。这里整表建网
            # 一次后按列数切片，行内容与 row.cells 完全一致（同一网格实现）。
            try:
                grid_cells = table._cells
                column_count = table._column_count
            except Exception:
                grid_cells = None
                column_count = 0
            for r_idx, row in enumerate(table.rows):
                if grid_cells is not None and column_count > 0:
                    row_cells = grid_cells[r_idx * column_count:(r_idx + 1) * column_count]
                else:
                    row_cells = row.cells
                cells_in_row = []
                c_idx = 0
                prev_tc = None
                for cell in row_cells:
                    tc = cell._tc
                    colspan = 1
                    try:
                        grid_span = tc.get_or_add_tcPr().find(f'{{{NS["w"]}}}gridSpan')
                        if grid_span is not None:
                            colspan = int(grid_span.get(f'{{{NS["w"]}}}val', '1'))
                    except Exception:
                        pass
                        
                    vmerge_val = 'none'
                    try:
                        vmerge = tc.get_or_add_tcPr().find(f'{{{NS["w"]}}}vMerge')
                        if vmerge is not None:
                            val = vmerge.get(f'{{{NS["w"]}}}val')
                            vmerge_val = val if val else 'continue'
                    except Exception:
                        pass
                    
                    if tc is prev_tc:
                        # 跳过合并单元格的重复引用文本叠加，但是为了对齐我们放入空字串
                        cells_in_row.append("")
                    else:
                        prev_tc = tc
                        rich_parts = []
                        for cell_paragraph in cell.paragraphs:
                            rendered = self._extract_rich_text(cell_paragraph).strip()
                            if rendered:
                                rich_parts.append(rendered)
                        text = "\n".join(rich_parts).strip() if rich_parts else cell.text.strip()
                        cells_in_row.append(text)
                        
                        # 仅对主单元格录入 cell_info
                        cells_info.append({
                            "row": r_idx,
                            "col": c_idx,
                            "text": text,
                            "row_span": 1,
                            "col_span": colspan,
                            "vmerge": vmerge_val
                        })
                    c_idx += 1
                rows.append(cells_in_row)
            
            if not rows:
                return None
            
            markdown_table = self._table_to_markdown(rows)
            return Block(
                id=block_id,
                type="table",
                text=markdown_table,
                metadata={
                    "source": "docx",
                    "cells": cells_info,
                    "artifact": False,
                    "in_table": True,
                    "canonical": True,
                },
                table_data={"rows": rows},
            )
        except Exception as e:
            logger.warning(f"处理表格时出错: {str(e)}")
            return None
    
    def _table_to_markdown(self, rows: List[List[str]]) -> str:
        """将表格数据渲染为 GFM 表格。

        委托全管道唯一入口 :meth:`Block.render_markdown_table`：按最大列宽对齐、
        **超宽行不截断**（旧实现以表头列数截断会丢列）、单元格内换行 → ``<br>``、
        竖线转义。
        """
        return Block.render_markdown_table(rows)
    
    def _extract_images_from_paragraph(
        self, paragraph: Paragraph, start_block_id: int
    ) -> List[Block]:
        """从段落中提取图片"""
        image_blocks = []
        block_id = start_block_id

        # 常见情形是段落根本没有图片；两次子树查找即可跳过逐 run 扫描。
        if not self._paragraph_has_inline_image(paragraph._element):
            return image_blocks

        try:
            for run in paragraph.runs:
                if hasattr(run, '_element'):
                    ns_w = f'{{{NS["w"]}}}'
                    for drawing in run._element.findall(f'.//{ns_w}drawing'):
                        try:
                            img_block = self._extract_image_from_drawing(drawing, block_id)
                            if img_block:
                                image_blocks.append(img_block)
                                block_id += 1
                        except Exception as e:
                            logger.warning(f"提取图片时出错: {str(e)}")
                    
                    for pict in run._element.findall(f'.//{ns_w}pict'):
                        try:
                            img_block = self._extract_image_from_pict(pict, block_id)
                            if img_block:
                                image_blocks.append(img_block)
                                block_id += 1
                        except Exception as e:
                            logger.warning(f"提取 VML 图片时出错: {str(e)}")
        except Exception as e:
            logger.warning(f"从段落提取图片时出错: {str(e)}")
        
        return image_blocks
    
    def _extract_image_from_drawing(self, drawing, block_id: int) -> Optional[Block]:
        """从 drawing 元素提取图片"""
        try:
            ns_a = f'{{{NS["a"]}}}'
            ns_r = f'{{{NS["r"]}}}'
            
            blip = drawing.find(f'.//{ns_a}blip')
            if blip is None:
                return None
            
            # Word 2016+ 在插入 SVG 时，会将真正的 SVG 存放在 a:extLst 中，
            # 而把自动生成的 PNG 后备图存放在 blip 的默认 r:embed 中。
            # 为了能够真正显示 SVG，我们必须优先尝试获取 SVG 的 embed_id：
            embed_id = None
            for ext in blip.findall(f'.//{ns_a}ext'):
                asvg_blip = ext.find('.//{http://schemas.microsoft.com/office/drawing/2016/SVG/main}svgBlip')
                if asvg_blip is not None:
                    embed_id = asvg_blip.get(f'{ns_r}embed')
                    if embed_id:
                        break
            
            # 如果没有找到 SVG，或者非 SVG 格式，则回退读取默认图片资源
            if not embed_id:
                embed_id = blip.get(f'{ns_r}embed')
                
            if not embed_id:
                return None
            
            image_uuid = f"img_{self.image_counter}_{hashlib.md5(embed_id.encode()).hexdigest()[:8]}"
            self.image_counter += 1
            
            image_data_str = f"[IMAGE_PLACEHOLDER: {image_uuid}]"
            if self._doc_rels and embed_id in self._doc_rels:
                try:
                    rel = self._doc_rels[embed_id]
                    if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'blob'):
                        blob = rel.target_part.blob
                        self.image_store[image_uuid] = blob
                        import base64
                        content_type = getattr(rel.target_part, 'content_type', 'image/png')
                        b64_data = base64.b64encode(blob).decode('utf-8')
                        image_data_str = f"data:{content_type};base64,{b64_data}"
                except Exception as e:
                    logger.debug(f"提取图片实体数据失败，使用占位符: {str(e)}")
            
            caption = ""
            ns_wp = f'{{{NS["wp"]}}}'
            desc_elem = drawing.find(f'.//{ns_wp}docPr')
            if desc_elem is None:
                desc_elem = drawing.find(f'.//{ns_a}docPr')
            if desc_elem is not None:
                caption = desc_elem.get('descr', '') or desc_elem.get('name', '')
            
            return Block(
                id=block_id,
                type="image",
                image_data=image_data_str,
                caption=caption,
                metadata={"embed_id": embed_id, "uuid": image_uuid},
            )
        except Exception as e:
            logger.warning(f"提取图片详情时出错: {str(e)}")
            return None
    
    def _extract_image_from_pict(self, pict, block_id: int) -> Optional[Block]:
        """从旧版 VML pict 元素提取图片"""
        try:
            ns_r = f'{{{NS["r"]}}}'
            ns_v = f'{{{NS["v"]}}}'
            
            imagedata = pict.find(f'.//{ns_v}imagedata')
            if imagedata is None:
                return None
            
            embed_id = imagedata.get(f'{ns_r}id')
            if not embed_id:
                return None
            
            image_uuid = f"img_{self.image_counter}_{hashlib.md5(embed_id.encode()).hexdigest()[:8]}"
            self.image_counter += 1
            
            image_data_str = f"[IMAGE_PLACEHOLDER: {image_uuid}]"
            if self._doc_rels and embed_id in self._doc_rels:
                try:
                    rel = self._doc_rels[embed_id]
                    if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'blob'):
                        blob = rel.target_part.blob
                        self.image_store[image_uuid] = blob
                        import base64
                        content_type = getattr(rel.target_part, 'content_type', 'image/png')
                        b64_data = base64.b64encode(blob).decode('utf-8')
                        image_data_str = f"data:{content_type};base64,{b64_data}"
                except Exception as e:
                    logger.debug(f"提取 VML 图片实体数据失败: {str(e)}")
            
            title = imagedata.get('title', '')
            return Block(
                id=block_id,
                type="image",
                image_data=image_data_str,
                caption=title,
                metadata={"embed_id": embed_id, "uuid": image_uuid},
            )
        except Exception as e:
            logger.warning(f"提取 VML 图片详情时出错: {str(e)}")
            return None
    
    def get_image_data(self, uuid: str) -> Optional[bytes]:
        """获取图片的二进制数据"""
        return self.image_store.get(uuid)
    
    def _log_debug_info(self, blocks: List[Block]):
        """输出调试信息"""
        logger.info("=" * 80)
        logger.info(f"前 10 个 Block 的详细信息（共 {len(blocks)} 个）：")
        for block in blocks[:10]:
            style = block.metadata.get("style", "N/A") if block.metadata else "N/A"
            source = block.metadata.get("source", "") if block.metadata else ""
            text_preview = (block.text[:50] + "...") if block.text and len(block.text) > 50 else (block.text or "N/A")
            features = []
            if block.is_heading_style:
                features.append(f"H{block.heading_level}")
            if block.is_bold:
                features.append("Bold")
            if block.font_size:
                features.append(f"Size:{block.font_size:.0f}")
            if block.alignment:
                features.append(f"Align:{block.alignment}")
            if source:
                features.append(f"Src:{source}")
            feature_str = f" [{', '.join(features)}]" if features else ""
            logger.info(f"  Block {block.id}: type={block.type}, style={style}{feature_str}, text={text_preview}")
        logger.info("=" * 80)
