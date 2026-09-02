from io import BytesIO
import base64
import sys
from pathlib import Path
import zipfile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent.parent))

from app.core.exceptions import ProviderError
from infrastructure.providers.docx_provider import DocxProvider


PNG_BYTES = (
    base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )
)


def _build_docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(buffer)
    return buffer.getvalue()


def _save_docx_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _rewrite_docx_entry(file_bytes: bytes, entry_name: str, old: bytes, new: bytes) -> bytes:
    output = BytesIO()
    with zipfile.ZipFile(BytesIO(file_bytes), "r") as input_zip, zipfile.ZipFile(output, "w") as output_zip:
        for zip_info in input_zip.infolist():
            data = input_zip.read(zip_info.filename)
            if zip_info.filename == entry_name:
                data = data.replace(old, new)
            output_zip.writestr(zip_info, data)
    return output.getvalue()


def _insert_inline_omml(file_bytes: bytes) -> bytes:
    math_xml = (
        b'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        b'<m:r><m:t>O(N)</m:t></m:r>'
        b'</m:oMath>'
    )
    return _rewrite_docx_entry(
        file_bytes,
        "word/document.xml",
        b"<w:t>prefix  suffix</w:t>",
        b"<w:t>prefix </w:t></w:r>" + math_xml + b"<w:r><w:t> suffix</w:t>",
    )


def test_docx_provider_normalizes_strict_ooxml_relationships():
    original = _build_docx_bytes("Hello strict")
    strict_like = _rewrite_docx_entry(
        original,
        "_rels/.rels",
        b"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument",
        b"http://purl.oclc.org/ooxml/officeDocument/relationships/officeDocument",
    )

    blocks = DocxProvider().extract_from_bytes(strict_like)

    assert any(block.text == "Hello strict" for block in blocks)


def test_docx_provider_rejects_invalid_zip_payloads():
    with pytest.raises(ProviderError) as exc_info:
        DocxProvider().extract_from_bytes(b"not-a-docx")

    assert "压缩包" in exc_info.value.message


def test_docx_provider_keeps_inline_omml_in_paragraph_order():
    original = _build_docx_bytes("prefix  suffix")
    with_inline_math = _insert_inline_omml(original)

    blocks = DocxProvider().extract_from_bytes(with_inline_math)

    assert blocks[0].type == "text"
    assert blocks[0].text == "prefix $O(N)$ suffix"
    assert not any(block.type == "formula" for block in blocks)


def test_docx_provider_renders_monospace_run_as_inline_code():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Use ")
    code_run = paragraph.add_run("foo_bar")
    code_run.font.name = "Consolas"
    paragraph.add_run(" here")

    blocks = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))

    assert blocks[0].type == "text"
    assert blocks[0].text == "Use `foo_bar` here"


def test_docx_provider_detects_single_line_code_block():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("print('hello')")
    run.font.name = "Consolas"

    blocks = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))

    assert blocks[0].type == "code"
    assert blocks[0].text == "print('hello')"


def test_docx_provider_preserves_center_alignment_on_single_line_code_block():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("const x = 1")
    run.font.name = "Consolas"

    blocks = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))

    assert blocks[0].type == "code"
    assert blocks[0].text == "const x = 1"
    assert blocks[0].alignment == "center"


def test_docx_provider_keeps_inline_image_in_paragraph_order():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("before ")
    image_run = paragraph.add_run()
    image_run.add_picture(BytesIO(PNG_BYTES))
    paragraph.add_run(" after")

    blocks = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))

    assert blocks[0].type == "text"
    assert blocks[0].text.startswith("before ![")
    assert blocks[0].text.endswith(" after")
    assert not any(block.type == "image" for block in blocks)


def test_docx_provider_keeps_image_only_paragraph_as_image_block():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(BytesIO(PNG_BYTES))

    blocks = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))

    assert any(block.type == "image" for block in blocks)


def test_docx_provider_skips_section_properties_noise():
    blocks = DocxProvider().extract_from_bytes(_build_docx_bytes("plain"))

    assert not any((block.text or "").startswith("[RAW_XML_NODE: sectPr]") for block in blocks)


def test_table_cell_newline_does_not_break_markdown_table():
    """表格单元格内换行会打断 GFM 表格 → 必须转 <br>（协作文档常见多行单元格）。"""
    rows = [
        ["功能", "编号", "改动点"],
        ["订单导出", "100034467\nExportOrderJob", "对接\n\n契约: a | b"],
    ]
    md = DocxProvider()._table_to_markdown(rows)
    lines = md.split("\n")
    # 表头 + 分隔 + 恰好 1 条数据行（单元格内换行没有另起表格行）
    assert len(lines) == 3
    assert lines[0] == "| 功能 | 编号 | 改动点 |"
    assert lines[1] == "| --- | --- | --- |"
    # 单元格内换行 → <br>，竖线 → 转义；数据行不含裸换行
    assert "100034467<br>ExportOrderJob" in lines[2]
    assert "对接<br><br>契约: a \\| b" in lines[2]
    assert "\n" not in lines[2]


def test_table_overwide_row_not_truncated():
    """数据行列数 > 表头时不得截断（合并单元格 / 不规则表格的真实场景）。"""
    rows = [["A", "B"], ["1", "2", "3"]]
    md = DocxProvider()._table_to_markdown(rows)
    lines = md.split("\n")
    assert lines[1] == "| --- | --- | --- |"   # 列宽按最大值对齐
    assert lines[2] == "| 1 | 2 | 3 |"           # 第 3 列未被丢弃


def _omml(xml: str):
    """把一段 OMML 片段解析成 lxml 元素（自动补 m: 命名空间）。"""
    pytest.importorskip("lxml")
    from lxml import etree

    wrapped = (
        '<m:root xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f"{xml}</m:root>"
    )
    return etree.fromstring(wrapped.encode("utf-8"))[0]


def test_omml_to_text_handles_common_structures():
    prov = DocxProvider()

    frac = _omml(
        "<m:f><m:num><m:r><m:t>a</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>b</m:t></m:r></m:den></m:f>"
    )
    assert prov._omml_to_text(frac) == "\\frac{a}{b}"

    ssup = _omml(
        "<m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e>"
        "<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>"
    )
    assert prov._omml_to_text(ssup) == "x^{2}"

    ssub = _omml(
        "<m:sSub><m:e><m:r><m:t>a</m:t></m:r></m:e>"
        "<m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>"
    )
    assert prov._omml_to_text(ssub) == "a_{i}"

    rad = _omml("<m:rad><m:deg/><m:e><m:r><m:t>y</m:t></m:r></m:e></m:rad>")
    assert prov._omml_to_text(rad) == "\\sqrt{y}"

    nary = _omml(
        '<m:nary><m:naryPr><m:chr m:val="∑"/></m:naryPr>'
        "<m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>"
        "<m:sup><m:r><m:t>n</m:t></m:r></m:sup>"
        "<m:e><m:r><m:t>i</m:t></m:r></m:e></m:nary>"
    )
    assert prov._omml_to_text(nary) == "∑_{i=1}^{n} i"

    delim = _omml(
        "<m:d><m:e><m:r><m:t>x</m:t></m:r></m:e></m:d>"
    )
    assert prov._omml_to_text(delim) == "(x)"


def test_docx_provider_emits_stable_run_atoms_without_changing_content():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Alpha ")
    bold_run = paragraph.add_run("Beta")
    bold_run.bold = True
    payload = _save_docx_bytes(doc)

    first = DocxProvider().extract_from_bytes(payload)[0]
    second = DocxProvider().extract_from_bytes(payload)[0]

    assert first.text == "Alpha **Beta**"
    assert first.to_markdown() == "Alpha **Beta**"
    atoms = first.metadata["atoms"]
    assert [atom["text"] for atom in atoms] == ["Alpha ", "Beta"]
    assert [atom["atom_id"] for atom in atoms] == [
        atom["atom_id"] for atom in second.metadata["atoms"]
    ]
    assert all(atom["block_id"] == first.id for atom in atoms)
    assert first.text[atoms[1]["char_start"]:atoms[1]["char_end"]] == "Beta"


def test_docx_provider_preserves_inherited_rich_run_markers_and_atom_styles():
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    character = doc.styles.add_style("Task Rich Strike", WD_STYLE_TYPE.CHARACTER)
    character.font.strike = True
    character.font.italic = True
    paragraph = doc.add_paragraph()
    paragraph.add_run("gone", style=character)
    paragraph.add_run(" ")
    superscript = paragraph.add_run("2")
    superscript.font.superscript = True

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.text == "*~~gone~~* <sup>2</sup>"
    atoms = block.metadata["atoms"]
    assert atoms[0]["is_strike"] is True
    assert atoms[0]["is_italic"] is True
    assert any(atom["is_superscript"] is True for atom in atoms)


def test_docx_provider_escapes_markdown_delimiters_in_source_text():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("literal ~~ and <tag>")
    run.font.underline = True

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.text == r"<u>literal \~\~ and \<tag\></u>"


def test_docx_provider_font_and_bold_are_character_weighted():
    from docx.shared import Pt

    doc = Document()
    paragraph = doc.add_paragraph()
    display = paragraph.add_run("X")
    display.bold = True
    display.font.size = Pt(24)
    body = paragraph.add_run(" ordinary body text dominates")
    body.font.size = Pt(10)

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.font_size == 10.0
    assert block.is_bold is False
    assert block.metadata["atoms"][0]["font_size"] == 24.0
    assert block.metadata["atoms"][1]["font_size"] == 10.0


def test_docx_provider_resolves_heading_and_font_through_based_on_chain():
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    doc = Document()
    base = doc.styles.add_style("Task5 Base", WD_STYLE_TYPE.PARAGRAPH)
    base.base_style = doc.styles["Heading 2"]
    base.font.bold = True
    base.font.size = Pt(17)
    derived = doc.styles.add_style("Task5 Derived", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base
    doc.add_paragraph("Inherited heading", style=derived)

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.is_heading_style is True
    assert block.heading_level == 2
    assert block.is_bold is True
    assert block.font_size == 17.0
    assert block.metadata["effective_style"] == "Heading 2"
    assert block.metadata["style_chain"][:2] == [derived.style_id, base.style_id]


def test_docx_provider_reads_style_numpr_ilvl():
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    numbered = doc.styles.add_style("Task5 Numbered", WD_STYLE_TYPE.PARAGRAPH)
    p_pr = numbered.element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "2")
    num_pr.append(ilvl)
    p_pr.append(num_pr)
    doc.add_paragraph("Styled list item", style=numbered)

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.metadata["list_level"] == 2
    assert block.metadata["numbering_level"] == 3
    assert block.has_heading_numbering is True


def test_docx_provider_inherits_numid_only_as_level_zero_and_honors_disable():
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    base = doc.styles.add_style("Task5 NumId Base", WD_STYLE_TYPE.PARAGRAPH)
    base_num_pr = OxmlElement("w:numPr")
    base_num_id = OxmlElement("w:numId")
    base_num_id.set(qn("w:val"), "7")
    base_num_pr.append(base_num_id)
    base.element.get_or_add_pPr().append(base_num_pr)
    derived = doc.styles.add_style("Task5 NumId Derived", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base

    doc.add_paragraph("Inherited numbered item", style=derived)
    disabled = doc.add_paragraph("Explicitly unnumbered item", style=derived)
    direct_num_pr = OxmlElement("w:numPr")
    direct_num_id = OxmlElement("w:numId")
    direct_num_id.set(qn("w:val"), "0")
    direct_num_pr.append(direct_num_id)
    disabled._p.get_or_add_pPr().append(direct_num_pr)

    inherited_block, disabled_block = DocxProvider().extract_from_bytes(
        _save_docx_bytes(doc)
    )

    assert inherited_block.metadata["list_level"] == 0
    assert inherited_block.metadata["numbering_level"] == 1
    assert inherited_block.has_heading_numbering is True
    assert disabled_block.metadata["list_level"] is None
    assert disabled_block.metadata["numbering_level"] is None
    assert disabled_block.has_heading_numbering is False


def test_docx_provider_dengxian_is_not_monospace():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("等线正文不是代码")
    run.font.name = "等线"

    block = DocxProvider().extract_from_bytes(_save_docx_bytes(doc))[0]

    assert block.type == "text"
    assert block.text == "等线正文不是代码"
    assert block.to_markdown() == "等线正文不是代码"


def test_docx_table_is_one_canonical_rendered_container():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A | B"
    table.cell(1, 1).text = "line 1\nline 2"

    table_block = next(
        block for block in DocxProvider().extract_from_bytes(_save_docx_bytes(doc))
        if block.type == "table"
    )

    assert table_block.metadata["source"] == "docx"
    assert table_block.metadata["canonical"] is True
    assert table_block.metadata["artifact"] is False
    assert table_block.metadata["in_table"] is True
    assert table_block.text == table_block.render_markdown_table(
        table_block.table_data["rows"]
    )
    assert "A \\| B" in table_block.text
    assert "line 1<br>line 2" in table_block.text


def test_docx_table_preserves_inline_rich_text_in_cells():
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Header"
    paragraph = table.cell(1, 0).paragraphs[0]
    paragraph.add_run("before ")
    strike = paragraph.add_run("removed")
    strike.font.strike = True
    paragraph.add_run(" ")
    bold = paragraph.add_run("important")
    bold.bold = True

    table_block = next(
        block for block in DocxProvider().extract_from_bytes(_save_docx_bytes(doc))
        if block.type == "table"
    )

    assert "before ~~removed~~ **important**" in table_block.table_data["rows"][1][0]
    assert "before ~~removed~~ **important**" in table_block.text
