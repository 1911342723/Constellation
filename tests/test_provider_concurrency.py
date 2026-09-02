"""Provider 并发隔离回归。

DocxProvider 携带 per-extraction 状态（image_store / _doc_rels / 样式链缓存），
registry 若返回共享单例，SSE 端点在 ``asyncio.to_thread`` 里并发抽块时，
一个请求的 ``_reset_state()`` 会清掉另一个请求正在用的文档上下文——静默地
把图片与样式串到别的文档。registry 必须 per-call 实例化。
"""
from __future__ import annotations

import base64
import io
from concurrent.futures import ThreadPoolExecutor

from docx import Document
from docx.shared import Pt

from infrastructure.providers import registry

# 1x1 红点 PNG
_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _docx_bytes(marker: str, *, bold: bool, with_image: bool) -> bytes:
    doc = Document()
    for index in range(30):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{marker} 段落 {index} 的内容")
        run.bold = bold
        run.font.size = Pt(14 if bold else 11)
    if with_image:
        doc.add_picture(io.BytesIO(_TINY_PNG))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _fingerprint(blocks) -> list[tuple]:
    return [(b.type, b.text, b.is_bold, b.font_size, b.image_data) for b in blocks]


def test_get_provider_returns_isolated_instances():
    assert registry.get_provider("docx") is not registry.get_provider("docx")


def test_concurrent_docx_extraction_matches_serial_reference():
    doc_a = _docx_bytes("A", bold=True, with_image=True)
    doc_b = _docx_bytes("B", bold=False, with_image=False)
    ref_a = _fingerprint(registry.extract_blocks("docx", doc_a))
    ref_b = _fingerprint(registry.extract_blocks("docx", doc_b))
    assert ref_a != ref_b  # 两份文档必须可区分，否则测试无效

    def run(job: tuple[bytes, list[tuple]]) -> bool:
        payload, reference = job
        return _fingerprint(registry.extract_blocks("docx", payload)) == reference

    jobs = [(doc_a, ref_a), (doc_b, ref_b)] * 8
    with ThreadPoolExecutor(max_workers=8) as pool:
        results = list(pool.map(run, jobs))
    assert all(results)
